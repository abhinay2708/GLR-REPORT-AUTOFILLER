import streamlit as st
import google.generativeai as genai
from docx import Document
import json
import os
import tempfile

# --- Utility Functions ---

def configure_gemini(api_key):
    """Configures the Gemini API with the provided key."""
    genai.configure(api_key=api_key)

def extract_data_from_reports(pdf_files, template_path):
    """
    Extracts data from PDF reports using Gemini to match the fields in the docx template.
    
    Args:
        pdf_files: List of uploaded PDF files (Streamlit UploadedFile objects).
        template_path: Path to the docx template.
        
    Returns:
        A dictionary containing the extracted key-value pairs.
    """
    
    uploaded_pdfs = []
    for pdf in pdf_files:
        # Save to temp file first because genai.upload_file needs a path
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(pdf.getvalue())
            tmp_pdf_path = tmp_pdf.name
        
        uploaded_file = genai.upload_file(tmp_pdf_path, mime_type="application/pdf")
        uploaded_pdfs.append(uploaded_file)
        os.remove(tmp_pdf_path) # Clean up local temp file

    doc = Document(template_path)
    template_text = "\n".join([p.text for p in doc.paragraphs])
    
    # Construct Prompt
    prompt = f"""
    You are an expert insurance adjuster assistant. 
    I have an insurance template (text provided below) and several photo reports (attached PDFs).
    
    Your task is to:
    1. Analyze the Template Text to identify the fields that need to be filled (e.g., Name, Date, Claim Number, Damage Description, etc.).
    2. Analyze the Photo Reports to find the corresponding information.
    3. Return a JSON object where the keys are the exact placeholders or field names found in the template, and the values are the extracted information from the reports.
    
    Template Text:
    {template_text}
    
    Output JSON only. Do not include markdown formatting like ```json ... ```.
    """
    
    # Using gemini-1.5-flash as it is a stable model for this task.
    model = genai.GenerativeModel("gemini-2.5-flash")
    
    # Generate content
    response = model.generate_content([prompt] + uploaded_pdfs)
    
    try:
        # clean potential markdown
        text = response.text.strip()
        if text.startswith("```json"):
            text = text[7:]
        if text.endswith("```"):
            text = text[:-3]
        return json.loads(text)
    except Exception as e:
        st.error(f"Error parsing JSON from Gemini: {e}")
        return {}

def fill_template(template_path, data, output_path):
    """
    Fills the docx template with the provided data.
    
    Args:
        template_path: Path to the input docx template.
        data: Dictionary of key-value pairs to fill.
        output_path: Path to save the filled docx.
    """
    doc = Document(template_path)
    
    # Simple text replacement in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
                
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
                            
    doc.save(output_path)

# --- Streamlit App ---

st.set_page_config(page_title="GLR Pipeline", layout="wide")

st.title("📄 GLR Pipeline: Insurance Template Automation")
st.markdown("""
Upload your Insurance Template (`.docx`) and Photo Reports (`.pdf`). 
The AI will extract data from the reports and fill the template for you.
""")

# Sidebar for Configuration
with st.sidebar:
    st.header("Configuration")
    api_key = st.text_input("Enter Gemini API Key", type="password")
    if not api_key:
        st.warning("Please enter your Gemini API Key to proceed.")

# Main Content
col1, col2 = st.columns(2)

with col1:
    st.subheader("1. Upload Documents")
    template_file = st.file_uploader("Upload Insurance Template (.docx)", type=["docx"])
    report_files = st.file_uploader("Upload Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("Generate Filled Document", type="primary", disabled=not (api_key and template_file and report_files)):
    if not api_key:
        st.error("API Key is missing!")
    else:
        configure_gemini(api_key)
        
        with st.spinner("Processing documents... This may take a moment."):
            try:
                # Save uploaded template to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_template:
                    tmp_template.write(template_file.getvalue())
                    tmp_template_path = tmp_template.name
                
                # Extract Data
                st.info("Extracting data from photo reports...")
                extracted_data = extract_data_from_reports(report_files, tmp_template_path)
                
                st.success("Data Extracted Successfully!")
                with st.expander("View Extracted Data"):
                    st.json(extracted_data)
                
                # Fill Template
                st.info("Filling template...")
                output_path = os.path.join(tempfile.gettempdir(), "filled_insurance_report.docx")
                fill_template(tmp_template_path, extracted_data, output_path)
                
                # Read result for download
                with open(output_path, "rb") as f:
                    file_data = f.read()
                
                st.success("Document Generated!")
                st.download_button(
                    label="Download Filled Document",
                    data=file_data,
                    file_name="filled_insurance_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Cleanup
                os.remove(tmp_template_path)
                
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
